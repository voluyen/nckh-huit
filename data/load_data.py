# ============================================================
# Unified Data Loader for Research-grade RAG Pipeline
# Support: PDF (text + OCR), DOCX, TXT/MD
# ============================================================

import os
from typing import List, Dict, Optional

import pdfplumber
import pytesseract
from pdf2image import convert_from_path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table


# ============================================================
# Helper: unified document schema
# ============================================================
def _make_doc(
    text: str,
    source: str,
    page: int,
    method: str,
    confidence: Optional[float] = None
) -> Dict:
    """
    method: text | ocr | docx | txt
    confidence: heuristic reliability score (for analysis, not ML)
    """
    return {
        "text": text,
        "source": source,
        "page": page,
        "method": method,
        "confidence": confidence
    }


# ============================================================
# TXT / MD loader
# ============================================================
def load_txt(path: str) -> List[Dict]:
    source = os.path.basename(path)

    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
    except Exception:
        return []

    if not text.strip():
        return []

    return [_make_doc(text, source, 1, "txt", 1.0)]


# ============================================================
# PDF loader (Text → OCR fallback, fully defensive)
# ============================================================
def load_pdf(path: str) -> List[Dict]:
    documents = []
    source = os.path.basename(path)

    try:
        pdf = pdfplumber.open(path)
    except Exception:
        return documents

    with pdf:
        for page_idx, page in enumerate(pdf.pages):
            page_no = page_idx + 1

            # ---------- STEP 1: Extract tables ----------
            table_texts = []
            try:
                tables = page.extract_tables()
            except Exception:
                tables = []

            for table in tables:
                if not table:
                    continue

                header = table[0]
                separator = ["---"] * len(header)

                md_table = "\n".join([
                    "| " + " | ".join(
                        (str(cell).strip() if cell else "")
                        for cell in row
                    ) + " |"
                    for row in [header, separator] + table[1:]
                ])
                table_texts.append("\n" + md_table + "\n")

            # ---------- STEP 2: Extract text ----------
            try:
                text = page.extract_text(layout=True)
            except Exception:
                text = None

            # ---------- STEP 3: OCR fallback ----------
            if not text or len(text.strip()) < 50:
                try:
                    images = convert_from_path(
                        path,
                        dpi=300,
                        first_page=page_no,
                        last_page=page_no,
                        thread_count=2
                    )
                    ocr_text = pytesseract.image_to_string(
                        images[0],
                        lang="vie+eng",
                        config="--oem 3 --psm 3"
                    )
                    text = ocr_text
                except Exception:
                    text = None

            # ---------- STEP 4: Merge content ----------
            if not text:
                continue

            full_content = text

            if table_texts:
                full_content += "\n\n### DATA TABLES\n" + "\n".join(table_texts)

            if len(full_content.strip()) >= 50:
                documents.append(
                    _make_doc(full_content, source, page_no, "text", 1.0)
                )

    return documents


# ============================================================
# DOCX loader (preserve reading order, paragraph + table)
# ============================================================
def load_docx(path: str) -> List[Dict]:
    source = os.path.basename(path)

    try:
        doc = Document(path)
    except Exception:
        return []

    blocks: List[str] = []

    # Duyệt theo thứ tự XML thật (không dùng doc.paragraphs/doc.tables)
    for child in doc.element.body:

        # ---------- Paragraph ----------
        if child.tag.endswith("p"):
            try:
                p = Paragraph(child, doc)
                if p.text and p.text.strip():
                    blocks.append(p.text.strip())
            except Exception:
                continue

        # ---------- Table ----------
        elif child.tag.endswith("tbl"):
            try:
                t = Table(child, doc)
                table_lines = []
                for row in t.rows:
                    row_text = "| " + " | ".join(
                        cell.text.strip().replace("\n", " ")
                        for cell in row.cells
                    ) + " |"
                    table_lines.append(row_text)

                if table_lines:
                    blocks.append("\n" + "\n".join(table_lines) + "\n")
            except Exception:
                continue

    if not blocks:
        return []

    return [
        _make_doc(
            "\n\n".join(blocks),
            source,
            1,
            "docx",
            1.0
        )
    ]


# ============================================================
# Unified entry point
# ============================================================
def load_data(path: str) -> List[Dict]:
    """
    Unified loader.
    Returns list of documents with metadata.
    """
    if not os.path.exists(path):
        raise FileNotFoundError(f"File không tồn tại: {path}")

    ext = os.path.splitext(path)[-1].lower()

    if ext == ".pdf":
        return load_pdf(path)
    elif ext == ".docx":
        return load_docx(path)
    elif ext in [".txt", ".md"]:
        return load_txt(path)
    else:
        raise ValueError(f"Định dạng {ext} không được hỗ trợ")
